from docx import Document
from docx.shared import Pt
import os
from reportlab.pdfgen import canvas

def generer_rapport(patient_nom, data, dossier_sortie, date_str):
    doc = Document()

    # Titre du rapport
    doc.add_heading(f"Rapport médical - {patient_nom}", level=1)

    for sheet, df in data.items():
        doc.add_heading(sheet, level=2)

        # Ajouter la table
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'

        # En-têtes
        hdr_cells = table.rows[0].cells
        for j, col in enumerate(df.columns):
            hdr_cells[j].text = str(col)

        # Lignes
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for j, value in enumerate(row):
                row_cells[j].text = str(value)

    # Sauvegarde Word
    nom_fichier = f"{patient_nom}_{date_str}.docx"
    chemin_word = os.path.join(dossier_sortie, nom_fichier)
    doc.save(chemin_word)

    # Export en PDF (via docx2pdf ou reportlab)
    try:
        from docx2pdf import convert
        convert(chemin_word, dossier_sortie)
    except:
        print("⚠️ docx2pdf non installé, PDF non généré.")

    return chemin_word
